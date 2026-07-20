from docx import Document as DocxDocument
from .base import BaseProcessor


class DocxProcessor(BaseProcessor):
    def extract_text(self, file_path: str) -> list[tuple[str, int | None]]:
        doc = DocxDocument(file_path)
        segments: list[tuple[str, int | None]] = []
        current_heading: str | None = None
        buffer: list[str] = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            if para.style.name.startswith("Heading"):
                if buffer:
                    segments.append(("\n".join(buffer), None))
                    buffer = []
                current_heading = para.text.strip()
                buffer.append(f"## {current_heading}")
            else:
                buffer.append(para.text.strip())

        if buffer:
            segments.append(("\n".join(buffer), None))

        return segments if segments else [("\n".join(p.text for p in doc.paragraphs if p.text.strip()), None)]

    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]
